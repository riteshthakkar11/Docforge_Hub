import html
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from backend.database import get_connection
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from docx import Document

router = APIRouter()


@router.get("/download/pdf/{document_id}")
def download_pdf(document_id: str):
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT dt.name FROM documents d
        JOIN document_templates dt ON d.template_id = dt.id
        WHERE d.id=%s
        """,
        (document_id,)
    )
    result = cursor.fetchone()
    if not result:
        raise HTTPException(status_code=404, detail="Document not found")
    title = result[0]

    cursor.execute(
        """
        SELECT DISTINCT ON (section_order)
            section_title, section_content, section_order
        FROM document_sections
        WHERE document_id=%s
        ORDER BY section_order, id DESC
        """,
        (document_id,)
    )
    sections = cursor.fetchall()
    if not sections:
        raise HTTPException(status_code=404, detail="No content found")

    file_name = title.lower().replace(" ", "_") + ".pdf"
    file_path = f"/tmp/{file_name}"

    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    title_style.alignment = TA_CENTER

    content = []
    content.append(Paragraph(html.escape(title), title_style))
    content.append(Spacer(1, 20))

    for row in sections:
        sec_title = row[0] or "Untitled Section"
        sec_text  = row[1] or "No content available"
        sec_text  = sec_text.replace("###", "")

        lines = sec_text.split("\n")
        if lines and lines[0].strip().lower() == sec_title.lower():
            lines = lines[1:]
        sec_text = "\n".join(lines).strip()
        sec_text = html.escape(sec_text)

        content.append(Paragraph(f"<b>{sec_title}</b>", styles["Heading2"]))
        content.append(Spacer(1, 10))
        for line in sec_text.split("\n"):
            if line.strip():
                content.append(Paragraph(line, styles["Normal"]))
                content.append(Spacer(1, 6))
        content.append(Spacer(1, 12))

    doc.build(content)
    cursor.close()
    conn.close()

    return FileResponse(
        path=file_path,
        filename=file_name,
        media_type="application/pdf"
    )


@router.get("/download/docx/{document_id}")
def download_docx(document_id: str):
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT section_title, section_content
        FROM document_sections
        WHERE document_id=%s
        ORDER BY section_order
        """,
        (document_id,)
    )
    sections = cursor.fetchall()

    doc = Document()
    for row in sections:
        sec_title = row[0] or "Untitled Section"
        sec_text  = row[1] or "No content available"
        sec_text  = sec_text.replace("###", "")

        lines = sec_text.split("\n")
        if lines and lines[0].strip().lower() == sec_title.lower():
            lines = lines[1:]
        sec_text = "\n".join(lines).strip()

        doc.add_heading(sec_title, level=1)
        doc.add_paragraph("")
        for line in sec_text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

    file_path = f"/tmp/{document_id}.docx"
    doc.save(file_path)

    cursor.close()
    conn.close()

    return FileResponse(
        path=file_path,
        filename=f"{document_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )